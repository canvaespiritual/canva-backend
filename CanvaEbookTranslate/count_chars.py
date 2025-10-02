from docx import Document

DOCX = r".\Alternativa de escrita ebook (1).docx"  # ajuste o nome se precisar

doc = Document(DOCX)
count = 0

for p in doc.paragraphs:
    count += len(p.text)

for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            count += len(c.text)

print("Caracteres aprox.:", count)
